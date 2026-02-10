"""
Document Builder — S1 Deterministic Report Renderer
Takes a ReportSchema + DataFrame and produces a polished Word document.
Uses docxtpl for template-based rendering, with a programmatic fallback.
"""

import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from capabilities.report_schema import ReportSchema


# ============================================================
# Color Palette (Professional Dark Blue Theme)
# ============================================================
HEADER_BG = RGBColor(0x1F, 0x3A, 0x5F)   # Dark navy blue
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)    # White text
ALT_ROW_BG = RGBColor(0xE8, 0xEE, 0xF4)  # Light blue-gray
ACCENT = RGBColor(0x2E, 0x75, 0xB6)       # Accent blue
TEXT_DARK = RGBColor(0x33, 0x33, 0x33)     # Soft black


def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)


def _format_header_cell(cell, text: str):
    """Format a table header cell with professional styling."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(9)
    run.font.color.rgb = HEADER_FG
    run.font.bold = True
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_shading(cell, "1F3A5F")


def _format_data_cell(cell, text: str, is_alt_row: bool = False):
    """Format a table data cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_DARK
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if is_alt_row:
        _set_cell_shading(cell, "E8EEF4")


def build_report(schema: ReportSchema, df: pd.DataFrame, output_path: str = None) -> str:
    """
    Build a professionally formatted Word document from a ReportSchema and DataFrame.
    
    Args:
        schema: The structured report data from the LLM
        df: The filtered DataFrame to include as a table
        output_path: Where to save the .docx file. Auto-generated if None.
        
    Returns:
        The absolute path to the generated .docx file
    """
    if output_path is None:
        # Generate output path based on title
        safe_title = "".join(c if c.isalnum() or c in (' ', '_') else '' for c in schema.title)
        safe_title = safe_title.strip().replace(' ', '_')[:50]
        
        # Save in user's Documents folder
        docs_dir = os.path.join(os.path.expanduser("~"), "Documents", "AI_Reports")
        os.makedirs(docs_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(docs_dir, f"{safe_title}_{timestamp}.docx")

    doc = Document()
    
    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)     # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Title ──
    title_para = doc.add_heading(schema.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title_para.runs:
        run.font.color.rgb = HEADER_BG
        run.font.size = Pt(26)

    # ── Metadata Line ──
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta_run.font.italic = True

    if schema.filter_description:
        filter_run = meta.add_run(f"  |  Filter: {schema.filter_description}")
        filter_run.font.size = Pt(9)
        filter_run.font.color.rgb = ACCENT
        filter_run.font.italic = True

    # ── Divider ──
    doc.add_paragraph("─" * 60)

    # ── Executive Summary ──
    summary_heading = doc.add_heading("Executive Summary", level=1)
    for run in summary_heading.runs:
        run.font.color.rgb = ACCENT
    
    summary_para = doc.add_paragraph(schema.summary)
    summary_para.paragraph_format.space_after = Pt(12)
    for run in summary_para.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK

    # ── Sections ──
    for section_data in schema.sections:
        heading = doc.add_heading(section_data.heading, level=2)
        for run in heading.runs:
            run.font.color.rgb = HEADER_BG

        para = doc.add_paragraph(section_data.content)
        para.paragraph_format.space_after = Pt(8)
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT_DARK

    # ── Data Table ──
    if schema.table_config.include and not df.empty:
        # Table caption
        caption = doc.add_heading(schema.table_config.caption, level=2)
        for run in caption.runs:
            run.font.color.rgb = HEADER_BG

        row_count_para = doc.add_paragraph()
        rc_run = row_count_para.add_run(f"Showing {len(df)} records")
        rc_run.font.size = Pt(9)
        rc_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        rc_run.font.italic = True

        # Limit to 100 rows for document sanity
        display_df = df.head(100)
        
        # Create table
        num_cols = len(display_df.columns)
        num_rows = len(display_df) + 1  # +1 for header
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header row
        for col_idx, col_name in enumerate(display_df.columns):
            _format_header_cell(table.rows[0].cells[col_idx], str(col_name))

        # Data rows
        for row_idx, (_, row) in enumerate(display_df.iterrows()):
            is_alt = row_idx % 2 == 1
            for col_idx, val in enumerate(row):
                _format_data_cell(
                    table.rows[row_idx + 1].cells[col_idx],
                    str(val) if pd.notna(val) else "",
                    is_alt_row=is_alt
                )

        if len(df) > 100:
            trunc_para = doc.add_paragraph()
            trunc_run = trunc_para.add_run(f"⚠ Showing first 100 of {len(df)} total records.")
            trunc_run.font.size = Pt(9)
            trunc_run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

    # ── Footer ──
    doc.add_paragraph("─" * 60)
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run("Generated by AI Agent — Document Intelligence Engine")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer_run.font.italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Save ──
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    return output_path
