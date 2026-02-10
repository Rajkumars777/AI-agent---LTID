"""
Document Operations - PDF, Word, and File Conversions
Combined module for all document-related operations.
"""

import pymupdf  # PyMuPDF
import pdfplumber  # S1-grade PDF table extraction
from docx import Document
import pandas as pd
from typing import List

# ==========================================
# PDF Operations
# ==========================================

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        doc = pymupdf.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"


def extract_tables_from_pdf(pdf_path: str) -> list:
    """
    Extract tables from a PDF using pdfplumber (S1-grade accuracy).
    Returns a list of tables, where each table is a list of rows (list of strings).
    """
    try:
        tables = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_table()
                if extracted:
                    tables.append(extracted)
        return tables if tables else [["No tables found in this PDF."]]
    except Exception as e:
        return [[f"Error extracting tables from PDF: {str(e)}"]]


# ==========================================
# Word Document Operations
# ==========================================

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a Word document."""
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"


def create_docx_with_text(output_path: str, text: str):
    """Create a new Word document with the given text."""
    try:
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)
        return f"Created {output_path}"
    except Exception as e:
        return f"Error creating DOCX: {str(e)}"


def replace_text_in_docx(docx_path: str, output_path: str, replacements: dict):
    """Replace text in a Word document based on a dictionary."""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)
        doc.save(output_path)
        return f"Saved modified doc to {output_path}"
    except Exception as e:
        return f"Error modifying DOCX: {str(e)}"


# ==========================================
# File Format Conversions
# ==========================================

def convert_csv_to_xlsx(csv_path: str, xlsx_path: str):
    """Convert CSV file to Excel XLSX format."""
    try:
        df = pd.read_csv(csv_path)
        df.to_excel(xlsx_path, index=False)
        return f"Converted {csv_path} to {xlsx_path}"
    except Exception as e:
        return f"Error converting CSV to XLSX: {str(e)}"


def convert_xlsx_to_csv(xlsx_path: str, csv_path: str):
    """Convert Excel XLSX file to CSV format."""
    try:
        df = pd.read_excel(xlsx_path)
        df.to_csv(csv_path, index=False)
        return f"Converted {xlsx_path} to {csv_path}"
    except Exception as e:
        return f"Error converting XLSX to CSV: {str(e)}"


def convert_json_to_csv(json_path: str, csv_path: str):
    """Convert JSON file to CSV format."""
    try:
        df = pd.read_json(json_path)
        df.to_csv(csv_path, index=False)
        return f"Converted {json_path} to {csv_path}"
    except Exception as e:
        return f"Error converting JSON to CSV: {str(e)}"
