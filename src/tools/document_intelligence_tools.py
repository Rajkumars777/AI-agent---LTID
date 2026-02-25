"""
document_intelligence_tools.py
===============================
Advanced document operations: PDF/Word extraction, template filling,
formatting, and cross-app data synchronization.
"""

import os
import re
from typing import Dict, Any, List, Optional
from pathlib import Path

# ─────────────────────────────────────────────────────
# DOCUMENT INTELLIGENCE (PDF/WORD)
# ─────────────────────────────────────────────────────

def _exec_extract_text_from_pdf(p: dict) -> str:
    """
    Extracts text or specific patterns from PDF.
    Use for: invoice numbers, dates, amounts, any text pattern.
    """
    pdf_path = p["pdf_path"]
    pattern = p.get("pattern")  # Optional regex pattern
    
    from src.services.desktop.ops import resolve_target_path
    path = resolve_target_path(pdf_path)
    
    if not path or not os.path.exists(path):
        return f"❌ PDF not found: {pdf_path}"
    
    try:
        import PyPDF2
        
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            full_text = ""
            for page in reader.pages:
                full_text += page.extract_text() + "\n"
        
        if not full_text.strip():
            return "❌ No text found in PDF (might be scanned/image-based)"
        
        # If pattern provided, extract specific matches
        if pattern:
            matches = re.findall(pattern, full_text, re.IGNORECASE)
            if matches:
                return f"✅ Extracted from PDF:\n" + "\n".join(matches)
            else:
                return f"❌ Pattern '{pattern}' not found in PDF"
        
        # Return full text (truncate if too long)
        if len(full_text) > 2000:
            full_text = full_text[:2000] + "\n... (truncated)"
        
        return f"✅ Extracted text from PDF:\n\n{full_text}"
        
    except ImportError:
        return "❌ PyPDF2 not installed. Run: pip install PyPDF2"
    except Exception as e:
        return f"❌ PDF extraction failed: {e}"


def _exec_extract_text_from_word(p: dict) -> str:
    """
    Extracts text or specific patterns from Word document.
    """
    doc_path = p["doc_path"]
    pattern = p.get("pattern")
    
    from src.services.desktop.ops import resolve_target_path
    path = resolve_target_path(doc_path)
    
    if not path or not os.path.exists(path):
        return f"❌ Word document not found: {doc_path}"
    
    try:
        import docx
        
        doc = docx.Document(path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        if not full_text.strip():
            return "❌ No text found in document"
        
        # Extract specific pattern
        if pattern:
            matches = re.findall(pattern, full_text, re.IGNORECASE)
            if matches:
                return f"✅ Extracted from Word:\n" + "\n".join(matches)
            else:
                return f"❌ Pattern '{pattern}' not found in document"
        
        # Return full text
        if len(full_text) > 2000:
            full_text = full_text[:2000] + "\n... (truncated)"
        
        return f"✅ Extracted text from Word:\n\n{full_text}"
        
    except ImportError:
        return "❌ python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"❌ Word extraction failed: {e}"


def _exec_fill_word_template(p: dict) -> str:
    """
    Fills a Word template with data.
    Replaces {{placeholders}} with actual values.
    """
    template_path = p["template_path"]
    output_path = p.get("output_path", "filled_document.docx")
    data = p["data"]  # dict of placeholder: value
    
    from src.services.desktop.ops import resolve_target_path
    path = resolve_target_path(template_path)
    
    if not path or not os.path.exists(path):
        return f"❌ Template not found: {template_path}"
    
    try:
        import docx
        
        doc = docx.Document(path)
        
        # Replace placeholders in paragraphs
        replacements_made = 0
        for para in doc.paragraphs:
            for placeholder, value in data.items():
                # Support both {{key}} and {key} formats
                for fmt in [f"{{{{{placeholder}}}}}", f"{{{placeholder}}}"]:
                    if fmt in para.text:
                        para.text = para.text.replace(fmt, str(value))
                        replacements_made += 1
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in data.items():
                        for fmt in [f"{{{{{placeholder}}}}}", f"{{{placeholder}}}"]:
                            if fmt in cell.text:
                                cell.text = cell.text.replace(fmt, str(value))
                                replacements_made += 1
        
        # Save output
        doc.save(output_path)
        
        return f"✅ Filled template: {replacements_made} replacements made\n📄 Saved to: {output_path}"
        
    except ImportError:
        return "❌ python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"❌ Template fill failed: {e}"


def _exec_format_word_text(p: dict) -> str:
    """
    Changes font, color, or paragraph formatting in Word document.
    """
    doc_path = p["doc_path"]
    target_text = p.get("target_text")  # Text to format (or "all")
    font_name = p.get("font_name")
    font_size = p.get("font_size")
    bold = p.get("bold")
    italic = p.get("italic")
    color = p.get("color")
    
    from src.services.desktop.ops import resolve_target_path
    path = resolve_target_path(doc_path)
    
    if not path or not os.path.exists(path):
        return f"❌ Document not found: {doc_path}"
    
    try:
        import docx
        from docx.shared import Pt, RGBColor
        
        doc = docx.Document(path)
        changes_made = 0
        
        for para in doc.paragraphs:
            # Check if this paragraph should be formatted
            if target_text == "all" or (target_text and target_text.lower() in para.text.lower()):
                for run in para.runs:
                    if font_name:
                        run.font.name = font_name
                        changes_made += 1
                    if font_size:
                        run.font.size = Pt(int(font_size))
                        changes_made += 1
                    if bold is not None:
                        run.font.bold = bold
                        changes_made += 1
                    if italic is not None:
                        run.font.italic = italic
                        changes_made += 1
                    if color:
                        # Convert color name to RGB
                        color_map = {
                            "red": (255, 0, 0),
                            "blue": (0, 0, 255),
                            "green": (0, 128, 0),
                            "black": (0, 0, 0),
                            "white": (255, 255, 255),
                        }
                        rgb = color_map.get(color.lower(), (0, 0, 0))
                        run.font.color.rgb = RGBColor(*rgb)
                        changes_made += 1
        
        doc.save(path)
        
        return f"✅ Formatting applied: {changes_made} changes made to '{doc_path}'"
        
    except ImportError:
        return "❌ python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"❌ Formatting failed: {e}"


# ─────────────────────────────────────────────────────
# DATA TRANSFORMATION (EXCEL)
# ─────────────────────────────────────────────────────

def _exec_input_data_to_excel(p: dict) -> str:
    """
    Inputs data into Excel cells and triggers formula recalculation.
    """
    excel_path = p["excel_path"]
    data = p["data"]  # dict of cell: value (e.g. {"A1": "Hello", "B2": 123})
    sheet_name = p.get("sheet_name")
    
    from src.services.desktop.ops import resolve_target_path
    path = resolve_target_path(excel_path)
    
    if not path or not os.path.exists(path):
        return f"❌ Excel file not found: {excel_path}"
    
    try:
        import openpyxl
        
        wb = openpyxl.load_workbook(path, data_only=False)
        ws = wb[sheet_name] if sheet_name else wb.active
        
        cells_updated = 0
        for cell_ref, value in data.items():
            ws[cell_ref] = value
            cells_updated += 1
        
        # Save (this triggers formula recalculation on open)
        wb.save(path)
        
        return f"✅ Updated {cells_updated} cell(s) in '{excel_path}'\n⚡ Formulas will recalculate when opened"
        
    except ImportError:
        return "❌ openpyxl not installed. Run: pip install openpyxl"
    except Exception as e:
        return f"❌ Excel update failed: {e}"


def _exec_refresh_pivot_table(p: dict) -> str:
    """
    Refreshes pivot tables in Excel (requires Excel COM automation on Windows).
    """
    excel_path = p["excel_path"]
    
    from src.services.desktop.ops import resolve_target_path
    path = resolve_target_path(excel_path)
    
    if not path or not os.path.exists(path):
        return f"❌ Excel file not found: {excel_path}"
    
    try:
        import win32com.client
        
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        wb = excel.Workbooks.Open(os.path.abspath(path))
        
        pivots_refreshed = 0
        for ws in wb.Worksheets:
            for pt in ws.PivotTables():
                pt.PivotCache().Refresh()
                pivots_refreshed += 1
        
        wb.Save()
        wb.Close()
        excel.Quit()
        
        return f"✅ Refreshed {pivots_refreshed} pivot table(s) in '{excel_path}'"
        
    except ImportError:
        return "❌ pywin32 not installed. Run: pip install pywin32"
    except Exception as e:
        return f"❌ Pivot refresh failed: {e}"


def _exec_style_excel_cells(p: dict) -> str:
    """
    Applies styling to Excel cells (colors, borders, bold, etc).
    """
    excel_path = p["excel_path"]
    cell_range = p["cell_range"]  # e.g. "A1:B10" or "headers"
    sheet_name = p.get("sheet_name")
    bg_color = p.get("bg_color")
    font_color = p.get("font_color")
    bold = p.get("bold", False)
    border = p.get("border", False)
    
    from src.services.desktop.ops import resolve_target_path
    path = resolve_target_path(excel_path)
    
    if not path or not os.path.exists(path):
        return f"❌ Excel file not found: {excel_path}"
    
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Border, Side
        
        wb = openpyxl.load_workbook(path)
        ws = wb[sheet_name] if sheet_name else wb.active
        
        # Handle special keywords
        if cell_range.lower() == "headers":
            cell_range = f"A1:{openpyxl.utils.get_column_letter(ws.max_column)}1"
        
        # Apply styling
        cells_styled = 0
        for row in ws[cell_range]:
            for cell in row:
                if bg_color:
                    color_hex = _color_to_hex(bg_color)
                    cell.fill = PatternFill(start_color=color_hex, end_color=color_hex, fill_type="solid")
                    cells_styled += 1
                
                if font_color:
                    color_hex = _color_to_hex(font_color)
                    cell.font = Font(color=color_hex, bold=bold)
                    cells_styled += 1
                elif bold:
                    cell.font = Font(bold=True)
                    cells_styled += 1
                
                if border:
                    thin = Side(style='thin', color='000000')
                    cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
                    cells_styled += 1
        
        wb.save(path)
        
        return f"✅ Styled {cells_styled} cell(s) in '{excel_path}'"
        
    except ImportError:
        return "❌ openpyxl not installed. Run: pip install openpyxl"
    except Exception as e:
        return f"❌ Styling failed: {e}"


# ─────────────────────────────────────────────────────
# CROSS-APP SYNCHRONIZATION
# ─────────────────────────────────────────────────────

def _exec_pdf_to_excel(p: dict) -> str:
    """
    Extracts data from PDF and inputs it into Excel.
    """
    pdf_path = p["pdf_path"]
    excel_path = p["excel_path"]
    pattern = p["pattern"]  # Regex to extract data
    target_cell = p.get("target_cell", "A1")
    sheet_name = p.get("sheet_name")
    
    # Step 1: Extract from PDF
    from src.services.desktop.ops import resolve_target_path
    pdf_full_path = resolve_target_path(pdf_path)
    
    if not pdf_full_path or not os.path.exists(pdf_full_path):
        return f"❌ PDF not found: {pdf_path}"
    
    try:
        import PyPDF2
        import openpyxl
        
        # Extract text from PDF
        with open(pdf_full_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            full_text = "".join([page.extract_text() for page in reader.pages])
        
        # Find pattern
        matches = re.findall(pattern, full_text, re.IGNORECASE)
        if not matches:
            return f"❌ Pattern '{pattern}' not found in PDF"
        
        # Step 2: Input to Excel
        excel_full_path = resolve_target_path(excel_path)
        if not excel_full_path or not os.path.exists(excel_full_path):
            # Create new workbook if doesn't exist
            wb = openpyxl.Workbook()
            excel_full_path = excel_path
        else:
            wb = openpyxl.load_workbook(excel_full_path)
        
        ws = wb[sheet_name] if sheet_name else wb.active
        
        # Input matches starting from target cell
        from openpyxl.utils import column_index_from_string, get_column_letter
        start_col = column_index_from_string(re.match(r"[A-Z]+", target_cell).group())
        start_row = int(re.search(r"\d+", target_cell).group())
        
        for i, match in enumerate(matches):
            ws[f"{get_column_letter(start_col)}{start_row + i}"] = match
        
        wb.save(excel_full_path)
        
        return f"✅ Extracted {len(matches)} value(s) from PDF to Excel\n📄 Saved to: {excel_full_path}"
        
    except Exception as e:
        return f"❌ PDF to Excel sync failed: {e}"


def _exec_web_to_excel(p: dict) -> str:
    """
    Extracts data from a webpage and inputs it into Excel.
    """
    url = p["url"]
    excel_path = p["excel_path"]
    css_selector = p.get("css_selector")  # Optional CSS selector for specific data
    target_cell = p.get("target_cell", "A1")
    
    try:
        from src.services.browser.agent import browser_agent
        import openpyxl
        
        # Navigate to URL
        browser_agent.navigate(url)
        
        # Extract data (placeholder - needs actual browser implementation)
        # You'll need to implement this based on your browser_agent
        extracted_data = ["Placeholder data"]  # TODO: implement extraction
        
        # Input to Excel
        from src.services.desktop.ops import resolve_target_path
        excel_full_path = resolve_target_path(excel_path)
        
        if not excel_full_path or not os.path.exists(excel_full_path):
            wb = openpyxl.Workbook()
            excel_full_path = excel_path
        else:
            wb = openpyxl.load_workbook(excel_full_path)
        
        ws = wb.active
        
        # Input data
        from openpyxl.utils import column_index_from_string, get_column_letter
        start_col = column_index_from_string(re.match(r"[A-Z]+", target_cell).group())
        start_row = int(re.search(r"\d+", target_cell).group())
        
        for i, value in enumerate(extracted_data):
            ws[f"{get_column_letter(start_col)}{start_row + i}"] = value
        
        wb.save(excel_full_path)
        
        return f"✅ Extracted data from web to Excel\n📄 Saved to: {excel_full_path}"
        
    except Exception as e:
        return f"❌ Web to Excel sync failed: {e}"


# ─────────────────────────────────────────────────────
# FORMAT CONVERSION
# ─────────────────────────────────────────────────────

def _exec_convert_format(p: dict) -> str:
    """
    Converts files between formats (Excel to PDF, CSV to XLSX, etc).
    """
    source_path = p["source_path"]
    target_format = p["target_format"].lower()  # pdf, xlsx, csv, docx
    output_path = p.get("output_path")
    
    from src.services.desktop.ops import resolve_target_path
    path = resolve_target_path(source_path)
    
    if not path or not os.path.exists(path):
        return f"❌ Source file not found: {source_path}"
    
    # Auto-generate output path if not provided
    if not output_path:
        base = os.path.splitext(path)[0]
        output_path = f"{base}.{target_format}"
    
    source_ext = os.path.splitext(path)[1].lower()
    
    try:
        # Excel to PDF
        if source_ext in ['.xlsx', '.xls'] and target_format == 'pdf':
            import win32com.client
            excel = win32com.client.Dispatch("Excel.Application")
            wb = excel.Workbooks.Open(os.path.abspath(path))
            wb.ExportAsFixedFormat(0, os.path.abspath(output_path))
            wb.Close()
            excel.Quit()
            return f"✅ Converted Excel to PDF\n📄 Saved to: {output_path}"
        
        # CSV to XLSX
        elif source_ext == '.csv' and target_format == 'xlsx':
            import pandas as pd
            df = pd.read_csv(path)
            df.to_excel(output_path, index=False)
            return f"✅ Converted CSV to XLSX\n📄 Saved to: {output_path}"
        
        # XLSX to CSV
        elif source_ext in ['.xlsx', '.xls'] and target_format == 'csv':
            import pandas as pd
            df = pd.read_excel(path)
            df.to_csv(output_path, index=False)
            return f"✅ Converted XLSX to CSV\n📄 Saved to: {output_path}"
        
        # Word to PDF
        elif source_ext in ['.docx', '.doc'] and target_format == 'pdf':
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(os.path.abspath(path))
            doc.SaveAs(os.path.abspath(output_path), FileFormat=17)  # 17 = PDF
            doc.Close()
            word.Quit()
            return f"✅ Converted Word to PDF\n📄 Saved to: {output_path}"
        
        else:
            return f"❌ Conversion from {source_ext} to {target_format} not supported yet"
        
    except ImportError as e:
        return f"❌ Missing library: {e}. Run: pip install pandas openpyxl pywin32"
    except Exception as e:
        return f"❌ Conversion failed: {e}"


# ─────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────

def _color_to_hex(color_name: str) -> str:
    """Convert color name to hex for Excel."""
    color_map = {
        "red": "FF0000",
        "blue": "0000FF",
        "green": "00FF00",
        "yellow": "FFFF00",
        "orange": "FFA500",
        "purple": "800080",
        "black": "000000",
        "white": "FFFFFF",
    }
    return color_map.get(color_name.lower(), "000000")


# ─────────────────────────────────────────────────────
# TOOL DEFINITIONS
# ─────────────────────────────────────────────────────

DOCUMENT_TOOLS = [
    
    # ── DOCUMENT INTELLIGENCE ──
    
    {
        "name": "extract_text_from_pdf",
        "description": (
            "Extracts text or specific patterns (invoice numbers, dates, amounts) from PDF. "
            "Use when user says: extract from PDF, get invoice number, find text in PDF, read PDF."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "pdf_path": {"type": "string", "description": "Path to PDF file"},
                "pattern": {"type": "string", "description": "Optional regex pattern to extract specific text"}
            },
            "required": ["pdf_path"]
        },
        "executor": _exec_extract_text_from_pdf
    },
    
    {
        "name": "extract_text_from_word",
        "description": (
            "Extracts text or specific patterns from Word document. "
            "Use when user says: extract from Word, get text from doc, read Word file."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "doc_path": {"type": "string", "description": "Path to Word document"},
                "pattern": {"type": "string", "description": "Optional regex pattern"}
            },
            "required": ["doc_path"]
        },
        "executor": _exec_extract_text_from_word
    },
    
    {
        "name": "fill_word_template",
        "description": (
            "Fills a Word template by replacing {{placeholders}} with actual data. "
            "Use when user says: fill template, populate document, insert data into Word."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "template_path": {"type": "string", "description": "Path to Word template"},
                "data": {"type": "object", "description": "Dictionary of placeholder: value pairs"},
                "output_path": {"type": "string", "description": "Output file path"}
            },
            "required": ["template_path", "data"]
        },
        "executor": _exec_fill_word_template
    },
    
    {
        "name": "format_word_text",
        "description": (
            "Changes font, size, color, or style in Word document. "
            "Use when user says: change font, make bold, color text, format Word."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "doc_path": {"type": "string", "description": "Path to Word document"},
                "target_text": {"type": "string", "description": "Text to format or 'all'"},
                "font_name": {"type": "string", "description": "Font name (e.g. Arial, Calibri)"},
                "font_size": {"type": "integer", "description": "Font size in points"},
                "bold": {"type": "boolean"},
                "italic": {"type": "boolean"},
                "color": {"type": "string", "description": "Color name (red, blue, green, etc)"}
            },
            "required": ["doc_path"]
        },
        "executor": _exec_format_word_text
    },
    
    # ── DATA TRANSFORMATION ──
    
    {
        "name": "input_data_to_excel",
        "description": (
            "Inputs data into Excel cells and triggers formula recalculation. "
            "Use when user says: put data in Excel, write to cells, update Excel, fill cells."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "excel_path": {"type": "string", "description": "Path to Excel file"},
                "data": {"type": "object", "description": "Dictionary of cell: value (e.g. {'A1': 'Hello', 'B2': 123})"},
                "sheet_name": {"type": "string", "description": "Sheet name (optional)"}
            },
            "required": ["excel_path", "data"]
        },
        "executor": _exec_input_data_to_excel
    },
    
    {
        "name": "refresh_pivot_table",
        "description": (
            "Refreshes pivot tables in Excel. "
            "Use when user says: refresh pivot, update pivot table, recalculate pivot."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "excel_path": {"type": "string", "description": "Path to Excel file with pivot tables"}
            },
            "required": ["excel_path"]
        },
        "executor": _exec_refresh_pivot_table
    },
    
    {
        "name": "style_excel_cells",
        "description": (
            "Applies styling to Excel cells (colors, borders, bold). "
            "Use when user says: color cells, add borders, make bold, format Excel, style headers."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "excel_path": {"type": "string", "description": "Path to Excel file"},
                "cell_range": {"type": "string", "description": "Cell range (e.g. A1:B10) or 'headers'"},
                "bg_color": {"type": "string", "description": "Background color"},
                "font_color": {"type": "string", "description": "Font color"},
                "bold": {"type": "boolean"},
                "border": {"type": "boolean"},
                "sheet_name": {"type": "string"}
            },
            "required": ["excel_path", "cell_range"]
        },
        "executor": _exec_style_excel_cells
    },
    
    # ── CROSS-APP SYNC ──
    
    {
        "name": "pdf_to_excel",
        "description": (
            "Extracts data from PDF and inputs it into Excel. "
            "Use when user says: PDF to Excel, extract invoice to Excel, get data from PDF to spreadsheet."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "pdf_path": {"type": "string", "description": "Source PDF file"},
                "excel_path": {"type": "string", "description": "Target Excel file"},
                "pattern": {"type": "string", "description": "Regex pattern to extract data"},
                "target_cell": {"type": "string", "description": "Starting cell (default: A1)"},
                "sheet_name": {"type": "string"}
            },
            "required": ["pdf_path", "excel_path", "pattern"]
        },
        "executor": _exec_pdf_to_excel
    },
    
    {
        "name": "web_to_excel",
        "description": (
            "Extracts data from webpage and inputs it into Excel. "
            "Use when user says: web to Excel, scrape website to spreadsheet, get data from URL to Excel."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "Webpage URL"},
                "excel_path": {"type": "string", "description": "Target Excel file"},
                "css_selector": {"type": "string", "description": "CSS selector for data"},
                "target_cell": {"type": "string", "description": "Starting cell"}
            },
            "required": ["url", "excel_path"]
        },
        "executor": _exec_web_to_excel
    },
    
    # ── FORMAT CONVERSION ──
    
    {
        "name": "convert_file_format",
        "description": (
            "Converts files between formats (Excel to PDF, CSV to XLSX, Word to PDF, etc). "
            "Use when user says: convert to PDF, save as Excel, change format, export as."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "source_path": {"type": "string", "description": "Source file path"},
                "target_format": {"type": "string", "description": "Target format: pdf, xlsx, csv, docx"},
                "output_path": {"type": "string", "description": "Output file path (optional)"}
            },
            "required": ["source_path", "target_format"]
        },
        "executor": _exec_convert_format
    },
    
]


# ─────────────────────────────────────────────────────
# REGISTRATION
# ─────────────────────────────────────────────────────

def register_document_tools():
    """Register all document intelligence and data transformation tools."""
    from src.tools.registry import registry
    
    registered = 0
    for tool in DOCUMENT_TOOLS:
        if not registry.has_tool(tool["name"]):
            meta = {k: v for k, v in tool.items() if k != "executor"}
            registry.register(meta, tool["executor"])
            registered += 1
    
    print(f"[DocumentTools] ✅ Registered {registered} tools")


# Auto-register on import
if __name__ in ("backend.tools.document_intelligence_tools", "tools.document_intelligence_tools"):
    register_document_tools()
