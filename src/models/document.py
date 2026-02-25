"""
Document Operations - PDF, Word, File Conversions, and Professional Report Building
Combined module for all document-related intelligence and operations.
"""

import os
import pymupdf  # PyMuPDF
import pdfplumber  # S1-grade PDF table extraction
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd
from typing import List, Optional
from datetime import datetime

# ==========================================
# PDF Operations
# ==========================================

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        doc = pymupdf.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

def extract_tables_from_pdf(pdf_path: str) -> list:
    """Extract tables from a PDF using pdfplumber."""
    try:
        tables = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_table()
                if extracted:
                    tables.append(extracted)
        return tables if tables else [["No tables found in this PDF."]]
    except Exception as e:
        return [[f"Error extracting tables from PDF: {str(e)}"]]

# ==========================================
# Word Document Operations
# ==========================================

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a Word document."""
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def create_docx_with_text(output_path: str, text: str):
    """Create a new Word document with the given text."""
    try:
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)
        return f"Created {output_path}"
    except Exception as e:
        return f"Error creating DOCX: {str(e)}"

def replace_text_in_docx(docx_path: str, output_path: str, replacements: dict):
    """Replace text in a Word document."""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)
        doc.save(output_path)
        return f"Saved modified doc to {output_path}"
    except Exception as e:
        return f"Error modifying DOCX: {str(e)}"

# ==========================================
# File Format Conversions
# ==========================================

def convert_csv_to_xlsx(csv_path: str, xlsx_path: str):
    """Convert CSV to XLSX."""
    df = pd.read_csv(csv_path)
    df.to_excel(xlsx_path, index=False)
    return f"Converted {csv_path} to {xlsx_path}"

def convert_xlsx_to_csv(xlsx_path: str, csv_path: str):
    """Convert XLSX to CSV."""
    df = pd.read_excel(xlsx_path)
    df.to_csv(csv_path, index=False)
    return f"Converted {xlsx_path} to {csv_path}"

# ==========================================
# S1 Report Builder (Professional)
# ==========================================

HEADER_BG = RGBColor(0x1F, 0x3A, 0x5F)
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_BG = RGBColor(0xE8, 0xEE, 0xF4)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
TEXT_DARK = RGBColor(0x33, 0x33, 0x33)

def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): color_hex, qn('w:val'): 'clear'})
    shading.append(shading_elm)

def _format_header_cell(cell, text: str):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size, run.font.bold, run.font.name = Pt(9), True, "Calibri"
    run.font.color.rgb = HEADER_FG
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_shading(cell, "1F3A5F")

def _format_data_cell(cell, text: str, is_alt_row: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size, run.font.name = Pt(9), "Calibri"
    run.font.color.rgb = TEXT_DARK
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if is_alt_row: _set_cell_shading(cell, "E8EEF4")

def build_report(schema, df: pd.DataFrame, output_path: str = None) -> str:
    """Build a professional report."""
    if output_path is None:
        safe_title = "".join(c if c.isalnum() or c in (' ', '_') else '' for c in schema.title).strip().replace(' ', '_')[:50]
        docs_dir = os.path.join(os.path.expanduser("~"), "Documents", "AI_Reports")
        os.makedirs(docs_dir, exist_ok=True)
        output_path = os.path.join(docs_dir, f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2)
    section.left_margin = section.right_margin = Cm(2.5)

    title_para = doc.add_heading(schema.title, level=0)
    for run in title_para.runs: run.font.color.rgb, run.font.size = HEADER_BG, Pt(26)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    meta_run.font.size, meta_run.font.color.rgb, meta_run.font.italic = Pt(9), RGBColor(0x88, 0x88, 0x88), True

    if schema.filter_description:
        filter_run = meta.add_run(f"  |  Filter: {schema.filter_description}")
        filter_run.font.size, filter_run.font.color.rgb, filter_run.font.italic = Pt(9), ACCENT, True

    doc.add_paragraph("─" * 60)
    summary_heading = doc.add_heading("Executive Summary", level=1)
    for run in summary_heading.runs: run.font.color.rgb = ACCENT
    summary_para = doc.add_paragraph(schema.summary)
    for run in summary_para.runs: run.font.size, run.font.color.rgb = Pt(11), TEXT_DARK

    for section_data in schema.sections:
        heading = doc.add_heading(section_data.heading, level=2)
        for run in heading.runs: run.font.color.rgb = HEADER_BG
        para = doc.add_paragraph(section_data.content)
        for run in para.runs: run.font.size, run.font.color.rgb = Pt(11), TEXT_DARK

    if schema.table_config.include and not df.empty:
        caption = doc.add_heading(schema.table_config.caption, level=2)
        for run in caption.runs: run.font.color.rgb = HEADER_BG
        display_df = df.head(100)
        table = doc.add_table(rows=len(display_df)+1, cols=len(display_df.columns))
        table.alignment, table.style = WD_TABLE_ALIGNMENT.CENTER, 'Table Grid'
        for col_idx, col_name in enumerate(display_df.columns):
            _format_header_cell(table.rows[0].cells[col_idx], str(col_name))
        for row_idx, (_, row) in enumerate(display_df.iterrows()):
            is_alt = row_idx % 2 == 1
            for col_idx, val in enumerate(row):
                _format_data_cell(table.rows[row_idx+1].cells[col_idx], str(val) if pd.notna(val) else "", is_alt_row=is_alt)

    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph("Generated by AI Agent — Document Intelligence Engine")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
